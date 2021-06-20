import re
from pathlib import Path
from typing import Optional

from docx import Document
from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger

from epc.processed.difficulties import difficulties
from epc.processed.ids import all_ids
from epc.processed.names_to_taxs import names_to_taxs
from epc.processed.reflection_tones import reflection_tones
from epc.processed.reflections import all_reflections
from epc.processed.taxonomies import all_taxonomies
from server.models import Passport

segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
morph_tagger = NewsMorphTagger(emb)

names_to_reviews = {}
for p_id, reflection in all_reflections.items():
    if p_id in all_ids:
        names_to_reviews[all_ids[p_id]] = reflection_tones[p_id]

excluded = [r"_Googl-Паспорт_образовательной_программы_pyVZaJh", "~$"]


def prepare_from(dir_path: Path, out_file: Path):
    comma = False
    with open(out_file, "w+") as out_f:
        out_f.write("[")
        docxs = dir_path.glob("*.docx")
        for f in docxs:
            file_name = f
            if any(e in str(file_name) for e in excluded):
                continue

            passport = read_docx(file_name)
            if not passport:
                continue

            if comma:
                out_f.write(",\n")
            comma = True

            out_f.write(repr(passport))
            out_f.flush()
        out_f.write("]")


def read_docx(file_name: Path) -> Optional[Passport]:
    document = Document(str(file_name))
    passport = Passport(str(file_name))
    try:
        for tbl in document.tables:
            for row in tbl.rows[1:]:
                if len(row.cells) < 3:
                    continue

                num_cell = row.cells[0].text.strip().lower()
                clean_nam = row.cells[1].text.strip().lower()
                clean_val = row.cells[2].text.strip().lower()

                if (
                    num_cell == "2.1" or "название" in clean_nam
                ) and not passport.prog_name:
                    doc = Doc(clean_val)
                    doc.segment(segmenter)
                    doc.tag_morph(morph_tagger)
                    for token in doc.tokens:
                        token.lemmatize(morph_vocab)
                    passport.prog_name = " ".join(t.lemma.lower() for t in doc.tokens)
                elif "час" in clean_nam and not passport.hours_cnt:
                    for g in re.findall(r"(\d+)", clean_val):
                        passport.hours_cnt += int(g)
                elif "формат обу" in clean_nam:
                    if "онлайн" in clean_val or "диста" in clean_val:
                        passport.online = True
                elif "уровень слож" in clean_nam and not passport.difficulty:
                    doc = Doc(clean_val.lower())
                    doc.segment(segmenter)
                    doc.tag_morph(morph_tagger)
                    for token in doc.tokens:
                        token.lemmatize(morph_vocab)
                    passport.difficulty = " ".join(
                        t.lemma for t in doc.tokens if t.lemma in difficulties.keys()
                    )
                elif "практикоор" in clean_nam and not passport.hours_prac:
                    for g in re.findall(r"(\d+)", clean_val):
                        h = int(g)
                        if passport.hours_prac + h > passport.hours_cnt:
                            break
                        passport.hours_prac += h
                elif (
                    "мини" in clean_nam
                    and "чело" in clean_nam
                    and not passport.min_listeners
                ):
                    for g in re.findall(r"(\d+)", clean_val):
                        passport.min_listeners += int(g)
                elif (
                    "макси" in clean_nam
                    and "чело" in clean_nam
                    and not passport.max_listeners
                ):
                    for g in re.findall(r"(\d+)", clean_val):
                        passport.max_listeners += int(g)

    except Exception as e:
        print(e)
        return None

    doc = Doc(" ".join(p.text.strip().lower() for p in document.paragraphs))
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    for token in doc.tokens:
        token.lemmatize(morph_vocab)
    doc_text = " ".join(t.lemma.lower() for t in doc.tokens)
    print(doc_text)

    tax_matches = []
    for tax_list in all_taxonomies:
        tax_match = 0
        for k, v in tax_list.items():
            if f" {k} " in doc_text:
                tax_match += 1
        tax_matches.append(tax_match)
    for k, v in all_taxonomies[tax_matches.index(max(tax_matches))].items():
        if f" {k} " in doc_text:
            passport.taxonomies.add(k)
            passport.taxonomy += v

    try:
        passport.expected_taxonomy_match = len(
            names_to_taxs.get(passport.prog_name, passport.taxonomies)
            & passport.taxonomies
        ) / len(passport.taxonomies)
    except:
        passport.expected_taxonomy_match = 1
    if not passport.prog_name:
        return None

    passport.reviews = names_to_reviews.get(passport.prog_name, passport.reviews)
    return passport
